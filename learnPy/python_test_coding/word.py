#!/usr/bin/python
# -*- coding: utf-8 -*-

import sys
from docx import Document
import time

def main(argv):
    print(argv[1])
    print(argv[2])
    top = int(argv[3])
    sec = int(argv[4])
    document = Document(argv[1]) 
    document.add_heading('Cloudscreen test microsoft office word docx', level=2)
    document.add_paragraph(time.asctime( time.localtime(time.time()) ))
    document.save(argv[1])
    for i in range(0, top):
        document = Document(argv[1])
    	document.add_heading('new paragraph {:d}'.format(i), level=3)
    	document.add_paragraph(argv[2])
    	print("docx modifed ............ {:d} ".format(i))
    	time.sleep(sec)
        document.save(argv[1])

if __name__ == '__main__':
    main(sys.argv)
